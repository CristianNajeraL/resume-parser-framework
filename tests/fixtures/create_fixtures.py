"""Script to create test fixture files."""

import io
from pathlib import Path

from docx import Document
from pypdf import PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def create_sample_pdf() -> None:
    """Create a sample PDF resume."""
    output_path = Path(__file__).parent / "sample_resume.pdf"

    # Create PDF with reportlab
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    # Add resume content
    c.drawString(100, 750, "Jane Doe")
    c.drawString(100, 730, "Email: jane.doe@example.com")
    c.drawString(100, 710, "")
    c.drawString(100, 690, "Skills:")
    c.drawString(100, 670, "- Python")
    c.drawString(100, 650, "- Machine Learning")
    c.drawString(100, 630, "- Data Science")

    c.save()

    # Save to file
    buffer.seek(0)
    with open(output_path, "wb") as f:
        f.write(buffer.read())

    print(f"Created {output_path}")


def create_sample_docx() -> None:
    """Create a sample Word resume."""
    output_path = Path(__file__).parent / "sample_resume.docx"

    doc = Document()
    doc.add_paragraph("John Smith")
    doc.add_paragraph("Email: john.smith@example.com")
    doc.add_paragraph("")
    doc.add_paragraph("Skills:")
    doc.add_paragraph("• Java")
    doc.add_paragraph("• Spring Boot")
    doc.add_paragraph("• Docker")

    doc.save(output_path)
    print(f"Created {output_path}")


def create_empty_pdf() -> None:
    """Create an empty PDF for testing."""
    output_path = Path(__file__).parent / "empty.pdf"

    writer = PdfWriter()
    with open(output_path, "wb") as f:
        writer.write(f)

    print(f"Created {output_path}")


if __name__ == "__main__":
    # Install required packages first:
    # uv add reportlab

    create_sample_pdf()
    create_sample_docx()
    create_empty_pdf()
    print("All fixtures created successfully!")
